from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path

from app.core.config import settings
from app.models.exam import Exam

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self) -> None:
        self.output_dir = Path(settings.UPLOAD_DIR) / "documents"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def generate_exam_document(
        self,
        exam: Exam,
        include_answer_key: bool = True,
    ) -> str:
        """Sınav için Word dokümanı üretir ve dosya yolunu döner."""
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        doc = docx.Document()

        # Sayfa kenar boşlukları
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        # Üst bilgi
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{exam.subject} | {exam.grade}. Sınıf | {exam.title}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(9)
        header_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Alt bilgi
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Oluşturulma Tarihi: {datetime.now().strftime('%d/%m/%Y')} | Sinavhazirlama"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(8)

        # Başlık
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(exam.title.upper())
        title_run.bold = True
        title_run.font.size = Pt(16)

        # Bilgi satırı
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = f"{exam.subject}  |  {exam.grade}. Sınıf"
        if exam.week_number:
            info_text += f"  |  {exam.week_number}. Hafta"
        info_para.add_run(info_text).font.size = Pt(11)

        doc.add_paragraph()

        # Ad/Soyad, Numara satırı
        meta_para = doc.add_paragraph()
        meta_para.add_run("Ad Soyad: ___________________________    ")
        meta_para.add_run("No: ____________    ")
        meta_para.add_run("Puan: ____________")

        doc.add_paragraph()

        # Sorular
        questions = sorted(exam.questions, key=lambda q: q.order_num)
        for i, question in enumerate(questions, start=1):
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{i}. ")
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_para.add_run(question.question_text).font.size = Pt(11)

            if question.question_type == "test" and question.options:
                try:
                    opts = json.loads(question.options)
                    for opt in opts:
                        opt_para = doc.add_paragraph(style="List Bullet")
                        opt_para.add_run(str(opt)).font.size = Pt(10)
                except (json.JSONDecodeError, TypeError):
                    pass
            elif question.question_type == "tf":
                tf_para = doc.add_paragraph()
                tf_para.add_run("   ( ) Doğru      ( ) Yanlış").font.size = Pt(10)
            elif question.question_type == "bosluk":
                pass  # Boşluk soru metninde zaten var
            elif question.question_type == "text":
                for _ in range(3):
                    line_para = doc.add_paragraph("_" * 70)
                    line_para.runs[0].font.size = Pt(10)

            doc.add_paragraph()

        # Cevap anahtarı
        if include_answer_key and questions:
            doc.add_page_break()
            key_title = doc.add_paragraph()
            key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            key_run = key_title.add_run("CEVAP ANAHTARI")
            key_run.bold = True
            key_run.font.size = Pt(14)
            doc.add_paragraph()

            for i, question in enumerate(questions, start=1):
                ans_para = doc.add_paragraph()
                ans_para.add_run(f"{i}. ").bold = True
                ans = question.correct_answer or "—"
                ans_para.add_run(f"{ans}").font.size = Pt(11)

        # Kaydet
        filename = f"sinav_{exam.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        logger.info("Sınav dokümanı oluşturuldu: %s", output_path)
        return str(output_path)
