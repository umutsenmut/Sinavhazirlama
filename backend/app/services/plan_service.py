from __future__ import annotations

import io
import logging
import os
from pathlib import Path

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.plan import Plan

logger = logging.getLogger(__name__)


class PlanService:
    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def upload_and_parse(
        self,
        content: bytes,
        filename: str,
        workspace_id: int,
    ) -> Plan:
        """Word dokümanını kaydeder ve kazanımları çıkarır."""
        upload_dir = Path(settings.UPLOAD_DIR) / str(workspace_id)
        upload_dir.mkdir(parents=True, exist_ok=True)

        safe_name = Path(filename).name
        file_path = upload_dir / safe_name
        file_path.write_bytes(content)

        title = Path(filename).stem
        plan = Plan(
            workspace_id=workspace_id,
            title=title,
            file_path=str(file_path),
            parsing_status="processing",
        )
        self.db.add(plan)
        await self.db.flush()
        await self.db.refresh(plan)

        try:
            parsed = await self._parse_docx(content)
            plan.parsed_content = parsed
            plan.parsing_status = "completed"
        except Exception as exc:
            logger.error("Plan ayrıştırma hatası (plan_id=%s): %s", plan.id, exc)
            plan.parsing_status = "failed"

        await self.db.flush()
        await self.db.refresh(plan)
        return plan

    async def _parse_docx(self, content: bytes) -> str:
        """Word dokümanından kazanım listesini çıkarır."""
        import docx

        doc = docx.Document(io.BytesIO(content))
        lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in lines:
                        lines.append(text)

        return "\n".join(lines)

    async def get_by_id(self, plan_id: int, workspace_id: int | None = None) -> Plan | None:
        query = select(Plan).where(Plan.id == plan_id)
        if workspace_id is not None:
            query = query.where(Plan.workspace_id == workspace_id)
        result = await self.db.execute(query)
        return result.scalar_one_or_none()

    async def list_by_workspace(
        self, workspace_id: int, page: int = 1, size: int = 20
    ) -> tuple[list[Plan], int]:
        count_result = await self.db.execute(
            select(func.count(Plan.id)).where(Plan.workspace_id == workspace_id)
        )
        total = count_result.scalar_one()

        result = await self.db.execute(
            select(Plan)
            .where(Plan.workspace_id == workspace_id)
            .order_by(Plan.created_at.desc())
            .offset((page - 1) * size)
            .limit(size)
        )
        return list(result.scalars().all()), total
